"""
utils/file_parser.py - Extract raw text from PDF and DOCX files.

Strategy:
  PDF  → pdfplumber (accurate text extraction) with PyMuPDF fallback
  DOCX → python-docx (paragraphs + tables)
"""

import re
from pathlib import Path
from backend.logger import logger


# ─── PDF ──────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract raw text from a PDF file using pdfplumber.
    Falls back to PyMuPDF (fitz) if pdfplumber fails.
    """
    path = Path(file_path)
    text = ""

    # Primary: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text:
                    pages.append(page_text)
            text = "\n".join(pages)
        if text.strip():
            logger.debug(f"PDF parsed via pdfplumber: {path.name} ({len(text)} chars)")
            return clean_text(text)
    except Exception as e:
        logger.warning(f"pdfplumber failed for {path.name}: {e}. Trying PyMuPDF...")

    # Fallback: PyMuPDF
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        pages = [doc[i].get_text("text") for i in range(len(doc))]
        text = "\n".join(pages)
        doc.close()
        if text.strip():
            logger.debug(f"PDF parsed via PyMuPDF: {path.name}")
            return clean_text(text)
    except Exception as e:
        logger.error(f"PyMuPDF also failed for {path.name}: {e}")

    raise ValueError(f"Could not extract text from PDF: {path.name}")


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def extract_text_from_docx(file_path: str) -> str:
    """
    Extract raw text from a DOCX file.
    Reads paragraphs and table cells.
    """
    from docx import Document

    path = Path(file_path)
    parts: list[str] = []

    try:
        doc = Document(str(path))

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        logger.debug(f"DOCX parsed: {path.name} ({len(text)} chars)")
        return clean_text(text)

    except Exception as e:
        logger.error(f"Failed to parse DOCX {path.name}: {e}")
        raise ValueError(f"Could not extract text from DOCX: {path.name}")


# ─── Router ───────────────────────────────────────────────────────────────────

def extract_text(file_path: str, file_type: str) -> str:
    """
    Route file to the correct parser based on extension.
    Returns cleaned plain text ready for AI processing.
    """
    file_type = file_type.lower().strip(".")
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ─── Text Cleaning ────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Normalize whitespace, remove junk characters, and standardize line breaks.
    """
    if not text:
        return ""

    # Normalize Unicode spaces and dashes
    text = text.replace("\u00a0", " ").replace("\u2013", "-").replace("\u2014", "-")

    # Remove null bytes and other control characters except newlines/tabs
    text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]", "", text)

    # Collapse multiple spaces (but preserve single newlines)
    text = re.sub(r"[ \t]+", " ", text)

    # Collapse 3+ consecutive newlines → 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines)

    return text.strip()
